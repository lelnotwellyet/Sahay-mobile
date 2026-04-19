import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_doc(doc_path, output_path):
    print(f"Opening {doc_path}...")
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return
    
    justified_count = 0
    font_changed = 0

    for paragraph in doc.paragraphs:
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            justified_count += 1
            
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            font_changed += 1
            
    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'

    try:
        doc.save(output_path)
        print(f"Saved formatted document to {output_path}. Justified {justified_count} paragraphs and set fonts to Times New Roman.")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == '__main__':
    doc_path = "/home/lelnotwellyet/Downloads/SAHAY_MAJOR (1).docx"
    output_path = "/home/lelnotwellyet/Downloads/SAHAY_MAJOR_Formatted.docx"
    format_doc(doc_path, output_path)
